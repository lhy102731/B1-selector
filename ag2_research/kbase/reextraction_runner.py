"""Recoverable, candidate-only runner for content re-extraction queues.

The runner extracts anchored text into candidate artifacts.  It never edits raw
objects, legacy packets, or the published catalog, and it never treats extracted
text (or an existing summary) as an accepted statement/evidence item.
"""
from __future__ import annotations

import argparse
from datetime import datetime, timezone
import hashlib
import json
import os
import re
from pathlib import Path
import shutil
import subprocess
import tempfile
from typing import Any, Callable

from .content_quality import validate_extraction_candidate
from .repository import KBaseRepository


RUN_SCHEMA_VERSION = 1
DEFAULT_RUN_ROOT = Path("wiki/outputs/candidates/ag2-kbase/content-layer-repair/reextraction-runs")
TESSERACT_EXE = Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe")


def _utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _atomic_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_name(path.name + ".tmp")
    temporary.write_text(json.dumps(value, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
                         encoding="utf-8")
    temporary.replace(path)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _artifact_stem(source_id: str, queue_id: str) -> str:
    queue_hash = hashlib.sha256(queue_id.encode("utf-8")).hexdigest()[:12]
    return f"{source_id}.{queue_hash}"


class _RunLock:
    """Best-effort single-writer guard for a queue run directory."""

    def __init__(self, path: Path) -> None:
        self.path = path
        self.fd: int | None = None

    def __enter__(self) -> "_RunLock":
        self.path.parent.mkdir(parents=True, exist_ok=True)
        try:
            self.fd = os.open(str(self.path), os.O_CREAT | os.O_EXCL | os.O_WRONLY)
            os.write(self.fd, str(os.getpid()).encode("ascii", errors="ignore"))
        except FileExistsError as error:
            raise RuntimeError(f"BLOCKED_RUN_ALREADY_ACTIVE: {self.path}") from error
        return self

    def __exit__(self, exc_type: Any, exc: Any, tb: Any) -> None:
        if self.fd is not None:
            os.close(self.fd)
            self.fd = None
        try:
            self.path.unlink()
        except FileNotFoundError:
            pass


def _segments_from_text(text: str, prefix: str = "line") -> list[dict[str, str]]:
    return [{"anchor": f"{prefix}:{number}", "text": line.strip()}
            for number, line in enumerate(text.splitlines(), 1) if line.strip()]


def _looks_like_good_text(text: str) -> bool:
    cleaned = re.sub(r"\s+", "", text or "")
    if len(cleaned) < 20:
        return False
    replacement_ratio = cleaned.count("\ufffd") / max(len(cleaned), 1)
    readable = sum(1 for char in cleaned if "\u4e00" <= char <= "\u9fff" or char.isalnum())
    readable_ratio = readable / max(len(cleaned), 1)
    return replacement_ratio < 0.03 and readable_ratio >= 0.35


def _blocked_ocr_message(kind: str, detail: str = "") -> RuntimeError:
    suggestion = (
        "BLOCKED_LOCAL_OCR_UNAVAILABLE: install Tesseract with Chinese language data "
        "(chi_sim+eng) or configure a local OCR worker; candidate-only extraction was not run"
    )
    if detail:
        suggestion = f"{suggestion}; {kind}: {detail}"
    return RuntimeError(suggestion)


def _render_page_to_pil(page: Any, dpi: int = 200) -> Any:
    import fitz  # type: ignore
    from PIL import Image  # type: ignore
    zoom = dpi / 72
    pixmap = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    return Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)


def _ocr_page_lines_with_pymupdf(page: Any, page_number: int, language: str) -> list[dict[str, str]]:
    if not hasattr(page, "get_textpage_ocr"):
        return []
    textpage = page.get_textpage_ocr(flags=0, language=language, dpi=200, full=True)
    return [{"anchor": f"page:{page_number}:ocr_line:{line_number}", "text": line.strip()}
            for line_number, line in enumerate(page.get_text("text", textpage=textpage).splitlines(), 1)
            if line.strip()]


def _ocr_page_lines_with_tesseract(page: Any, page_number: int, language: str) -> list[dict[str, str]]:
    import pytesseract  # type: ignore
    if TESSERACT_EXE.is_file():
        pytesseract.pytesseract.tesseract_cmd = str(TESSERACT_EXE)
    image = _render_page_to_pil(page)
    text = pytesseract.image_to_string(image, lang=language)
    return [{"anchor": f"page:{page_number}:ocr_line:{line_number}", "text": line.strip()}
            for line_number, line in enumerate(text.splitlines(), 1) if line.strip()]


def _extract_pdf_ocr(path: Path, *, language: str = "chi_sim+eng") -> tuple[list[dict[str, str]], str]:
    import fitz  # type: ignore
    errors: list[str] = []
    segments: list[dict[str, str]] = []
    with fitz.open(path) as document:
        for page_number, page in enumerate(document, 1):
            try:
                page_segments = _ocr_page_lines_with_pymupdf(page, page_number, language)
            except Exception as error:
                errors.append(f"pymupdf_ocr page {page_number}: {error}")
                page_segments = []
            if not page_segments:
                try:
                    page_segments = _ocr_page_lines_with_tesseract(page, page_number, language)
                except Exception as error:
                    errors.append(f"pytesseract page {page_number}: {error}")
                    page_segments = []
            segments.extend(page_segments)
    if not segments:
        raise _blocked_ocr_message("pdf", "; ".join(errors[-3:]))
    return segments, "local_ocr_pdf_pages"


def _extract_pdf(path: Path) -> tuple[list[dict[str, str]], str]:
    import fitz  # type: ignore
    segments: list[dict[str, str]] = []
    with fitz.open(path) as document:
        for page_number, page in enumerate(document, 1):
            for line_number, line in enumerate(page.get_text("text").splitlines(), 1):
                if line.strip():
                    segments.append({"anchor": f"page:{page_number}:line:{line_number}",
                                     "text": line.strip()})
    if segments and _looks_like_good_text("\n".join(segment["text"] for segment in segments[:200])):
        return segments, "pymupdf_native_text"
    return _extract_pdf_ocr(path)


def _extract_docx(path: Path) -> tuple[list[dict[str, str]], str]:
    from docx import Document  # type: ignore
    document = Document(path)
    segments = [{"anchor": f"paragraph:{number}", "text": paragraph.text.strip()}
                for number, paragraph in enumerate(document.paragraphs, 1)
                if paragraph.text.strip()]
    for table_number, table in enumerate(document.tables, 1):
        for row_number, row in enumerate(table.rows, 1):
            text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if text:
                segments.append({"anchor": f"table:{table_number}:row:{row_number}", "text": text})
    return segments, "python_docx_native_text"


def _extract_doc(path: Path) -> tuple[list[dict[str, str]], str]:
    soffice = shutil.which("soffice") or r"C:\Program Files\LibreOffice\program\soffice.exe"
    if not Path(soffice).is_file():
        raise RuntimeError("LibreOffice is unavailable for legacy .doc extraction")
    with tempfile.TemporaryDirectory(prefix="kbase-doc-extract-") as temporary:
        result = subprocess.run([soffice, "--headless", "--convert-to", "txt:Text", "--outdir",
                                 temporary, str(path)], capture_output=True, text=True, timeout=120,
                                check=False)
        output = Path(temporary) / f"{path.stem}.txt"
        if result.returncode or not output.is_file():
            raise RuntimeError(f"LibreOffice conversion failed: {result.stderr.strip()[:300]}")
        text = output.read_text(encoding="utf-8", errors="replace")
        if _looks_like_good_text(text):
            return _segments_from_text(text), "libreoffice_native_text"
        pdf_result = subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir",
                                     temporary, str(path)], capture_output=True, text=True,
                                    timeout=120, check=False)
        pdf_output = Path(temporary) / f"{path.stem}.pdf"
        if pdf_result.returncode or not pdf_output.is_file():
            raise RuntimeError(
                "BLOCKED_LEGACY_DOC_GARBLED_TEXT: LibreOffice text conversion was not readable "
                f"and PDF fallback failed: {pdf_result.stderr.strip()[:300]}"
            )
        return _extract_pdf_ocr(pdf_output)


def _extract_image_ocr(path: Path, *, language: str = "chi_sim+eng") -> tuple[list[dict[str, str]], str]:
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except Exception as error:
        raise _blocked_ocr_message("image", str(error))
    if TESSERACT_EXE.is_file():
        pytesseract.pytesseract.tesseract_cmd = str(TESSERACT_EXE)
    try:
        text = pytesseract.image_to_string(Image.open(path), lang=language)
    except Exception as error:
        raise _blocked_ocr_message("image", str(error))
    segments = _segments_from_text(text, prefix="image_ocr_line")
    if not segments:
        raise RuntimeError("BLOCKED_EMPTY_OR_IMAGE_ONLY_EXTRACTION")
    return segments, "local_ocr_image"


def extract_local(path: Path, task: dict[str, Any]) -> tuple[list[dict[str, str]], str]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".doc":
        return _extract_doc(path)
    if suffix in {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp"}:
        return _extract_image_ocr(path)
    if suffix in {".txt", ".md"}:
        return _segments_from_text(path.read_text(encoding="utf-8", errors="replace")), "plain_text"
    if task.get("route") == "requires_asr":
        raise RuntimeError("BLOCKED_LOCAL_ASR_UNAVAILABLE")
    raise RuntimeError("BLOCKED_LOCAL_OCR_OR_EXTRACTOR_UNAVAILABLE")


def _unpack_extraction(value: Any) -> tuple[list[dict[str, str]], str, dict[str, Any] | None]:
    """Accept built-in output or a future plugin's optional distilled candidate."""
    if not isinstance(value, tuple) or len(value) not in {2, 3}:
        raise TypeError("extractor must return (segments, method[, candidate])")
    segments, method = value[0], value[1]
    candidate = value[2] if len(value) == 3 else None
    if not isinstance(segments, list) or not isinstance(method, str):
        raise TypeError("extractor segments must be a list and method must be text")
    if candidate is not None and not isinstance(candidate, dict):
        raise TypeError("extractor candidate must be an object")
    return segments, method, candidate


def _quality_decision(candidate: dict[str, Any] | None, *, source_id: str,
                      raw_path: str) -> tuple[dict[str, Any], dict[str, Any] | None]:
    """Fail closed: anchored text is not a publication candidate."""
    if candidate is None:
        # Calling the shared gate here is intentional: there is one gate for
        # every extraction output, including non-distilled anchored text.
        gate_input = {"source_id": source_id, "raw_path": raw_path}
        result = validate_extraction_candidate(gate_input)
        return ({**result, "publication_eligible": False,
                 "stage": "anchored_text_awaiting_distillation"}, None)
    gate_input = dict(candidate)
    provenance_errors: list[str] = []
    if gate_input.get("source_id") not in {None, "", source_id}:
        provenance_errors.append("candidate source_id does not match queue task")
    if gate_input.get("raw_path") not in {None, "", raw_path}:
        provenance_errors.append("candidate raw_path does not match queue task")
    gate_input.setdefault("source_id", source_id)
    gate_input.setdefault("raw_path", raw_path)
    result = validate_extraction_candidate(gate_input)
    if provenance_errors:
        result = {"decision": "reject", "errors": provenance_errors + result.get("errors", []),
                  "warnings": result.get("warnings", [])}
    result = {**result, "publication_eligible": result.get("decision") == "accept",
              "stage": "distilled_candidate_quality_gate"}
    return result, gate_input


def run_queue(*, queue_path: str | Path, vault_path: str | Path | None = None,
              output_root: str | Path | None = None, dry_run: bool = False,
              limit: int | None = None, retry_failed: bool = False,
              retry_blocked: bool = False,
              extractor: Callable[[Path, dict[str, Any]], Any] = extract_local,
              crash_hook: Callable[[dict[str, Any]], None] | None = None) -> dict[str, Any]:
    """Run pending tasks with atomic state, per-item isolation, and resume support."""
    repo = KBaseRepository(vault_path)
    queue_file = Path(queue_path).resolve()
    queue = json.loads(queue_file.read_text(encoding="utf-8"))
    tasks = queue.get("tasks")
    if not isinstance(tasks, list):
        raise ValueError("queue.tasks must be a list")
    queue_hash = _sha256(queue_file)
    root = (repo.vault / (output_root or DEFAULT_RUN_ROOT) / queue_hash[:16]).resolve()
    root.relative_to(repo.vault)
    state_path = root / "state.json"
    audit_path = root / "audit.jsonl"
    if state_path.is_file():
        state = json.loads(state_path.read_text(encoding="utf-8"))
        if state.get("queue_sha256") != queue_hash:
            raise ValueError("run state belongs to a different queue")
    else:
        state = {"run_schema_version": RUN_SCHEMA_VERSION, "queue_sha256": queue_hash,
                 "queue_path": str(queue_file), "created_at": _utc_now(), "items": {}}
    # A process death after the in-flight checkpoint is safely recoverable.
    for item in state["items"].values():
        if item.get("status") == "in_flight":
            item.update(status="pending", recovery="interrupted_in_flight")

    selected: list[dict[str, Any]] = []
    for task in tasks:
        queue_id = str(task.get("queue_id") or "")
        prior = state["items"].get(queue_id, {})
        if prior.get("status") == "done":
            continue
        if prior.get("status") == "blocked" and not retry_blocked:
            continue
        if prior.get("status") == "failed" and not retry_failed:
            continue
        selected.append(task)
        if limit is not None and len(selected) >= limit:
            break
    if dry_run:
        return {"dry_run": True, "queue_sha256": queue_hash, "would_attempt": len(selected),
                "remaining_total": len(tasks) - sum(
                    state["items"].get(str(t.get("queue_id")), {}).get("status") == "done" for t in tasks),
                "output_root": str(root)}

    root.mkdir(parents=True, exist_ok=True)
    _atomic_json(state_path, state)

    def audit(event: dict[str, Any]) -> None:
        with audit_path.open("a", encoding="utf-8") as handle:
            handle.write(json.dumps({"at": _utc_now(), **event}, ensure_ascii=False, sort_keys=True) + "\n")

    for task in selected:
        queue_id = str(task.get("queue_id") or "")
        source_id = str(task.get("source_id") or "")
        item = state["items"].setdefault(queue_id, {})
        item.update(status="in_flight", source_id=source_id, started_at=_utc_now(),
                    attempts=int(item.get("attempts", 0)) + 1)
        _atomic_json(state_path, state)
        audit({"event": "started", "queue_id": queue_id, "source_id": source_id})
        if crash_hook:
            crash_hook(task)
        try:
            raw_path = task.get("raw_path")
            if not raw_path:
                raise RuntimeError("BLOCKED_RAW_UNAVAILABLE")
            raw = repo.safe_path(str(raw_path))
            if not raw.is_file():
                raise RuntimeError("BLOCKED_RAW_UNAVAILABLE")
            segments, method, distilled_candidate = _unpack_extraction(extractor(raw, task))
            segments = [segment for segment in segments
                        if str(segment.get("anchor", "")).strip() and str(segment.get("text", "")).strip()]
            if not segments:
                raise RuntimeError("BLOCKED_EMPTY_OR_IMAGE_ONLY_EXTRACTION")
            quality_gate, gated_candidate = _quality_decision(
                distilled_candidate, source_id=source_id, raw_path=str(raw_path))
            artifact = {
                "artifact_schema_version": 1, "source_id": source_id, "queue_id": queue_id,
                "raw_path": str(raw_path), "raw_sha256": _sha256(raw), "extraction_method": method,
                "segments": segments, "quality_gate": quality_gate,
                "policy": {"candidate_only": True, "raw_modified": False,
                           "packet_modified": False, "catalog_modified": False,
                           "external_models_called": False,
                           "distilled_statements": gated_candidate is not None,
                           "accepted_evidence": quality_gate["publication_eligible"],
                           "publication_eligible": quality_gate["publication_eligible"],
                           "notice": "Extracted text is input for later review/distillation, not evidence by itself."},
            }
            if gated_candidate is not None:
                artifact["distilled_candidate"] = gated_candidate
            artifact_path = root / "artifacts" / f"{_artifact_stem(source_id, queue_id)}.extracted.json"
            _atomic_json(artifact_path, artifact)
            item.update(status="done", completed_at=_utc_now(), extraction_method=method,
                        segment_count=len(segments), artifact_path=str(artifact_path.relative_to(repo.vault)).replace("\\", "/"))
            item.update(quality_decision=quality_gate["decision"],
                        publication_eligible=quality_gate["publication_eligible"])
            audit({"event": "done", "queue_id": queue_id, "source_id": source_id,
                   "segment_count": len(segments), "method": method})
        except Exception as error:  # per-item failure isolation is deliberate
            message = f"{type(error).__name__}: {error}"
            status = "blocked" if "BLOCKED_" in str(error) else "failed"
            item.update(status=status, completed_at=_utc_now(), error=message[:1000])
            audit({"event": status, "queue_id": queue_id, "source_id": source_id,
                   "error": message[:1000]})
        _atomic_json(state_path, state)

    counts: dict[str, int] = {}
    for item in state["items"].values():
        counts[item["status"]] = counts.get(item["status"], 0) + 1
    return {"dry_run": False, "queue_sha256": queue_hash, "attempted": len(selected),
            "counts": counts, "state_path": str(state_path), "audit_path": str(audit_path)}


def main() -> None:
    parser = argparse.ArgumentParser(description="Run candidate-only KBase content re-extraction")
    parser.add_argument("queue")
    parser.add_argument("--vault", default=None)
    parser.add_argument("--output-root", default=None)
    parser.add_argument("--dry-run", action="store_true")
    parser.add_argument("--limit", type=int, default=None)
    parser.add_argument("--retry-failed", action="store_true")
    parser.add_argument("--retry-blocked", action="store_true")
    args = parser.parse_args()
    print(json.dumps(run_queue(queue_path=args.queue, vault_path=args.vault,
                               output_root=args.output_root, dry_run=args.dry_run,
                               limit=args.limit, retry_failed=args.retry_failed,
                               retry_blocked=args.retry_blocked),
                     ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
